import os
import markdown
from docx import Document
from fpdf import FPDF
import re

def read_md_file(file_path):
    """读取MD文件内容"""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def convert_md_to_docx(md_content, output_path):
    """将MD内容转换为DOCX格式"""
    # 创建Document对象
    doc = Document()
    
    # 简单的MD到DOCX转换
    lines = md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # 处理粗体文本
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('> '):
            # 处理引用
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[2:])
            run.italic = True
        elif '|' in line and line.startswith('|'):
            # 简单处理表格
            doc.add_paragraph(line)
        else:
            if line.strip():
                doc.add_paragraph(line)
    
    # 保存DOCX文件
    doc.save(output_path)
    print(f"DOCX文件已保存到: {output_path}")

def convert_md_to_pdf(md_content, output_path):
    """将MD内容转换为PDF格式"""
    # 创建FPDF对象
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # 添加中文字体支持
    try:
        pdf.add_font("SimHei", "", "simhei.ttf")
        pdf.set_font("SimHei", size=12)
    except:
        # 如果无法加载中文字体，使用默认字体
        pdf.set_font("Arial", size=12)
    
    # 简单的MD到PDF转换
    lines = md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, line[2:].encode('latin-1', 'replace').decode('latin-1'), ln=True)
            pdf.set_font("Arial", size=12)
        elif line.startswith('## '):
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, line[3:].encode('latin-1', 'replace').decode('latin-1'), ln=True)
            pdf.set_font("Arial", size=12)
        elif line.startswith('### '):
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, line[4:].encode('latin-1', 'replace').decode('latin-1'), ln=True)
            pdf.set_font("Arial", size=12)
        elif line.startswith('- '):
            pdf.cell(0, 10, ("• " + line[2:]).encode('latin-1', 'replace').decode('latin-1'), ln=True)
        elif line.startswith('> '):
            pdf.set_font("Arial", "I", 12)
            pdf.cell(0, 10, line[2:].encode('latin-1', 'replace').decode('latin-1'), ln=True)
            pdf.set_font("Arial", size=12)
        else:
            if line.strip():
                pdf.cell(0, 10, line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
    
    # 保存PDF文件
    pdf.output(output_path)
    print(f"PDF文件已保存到: {output_path}")

def main():
    # 定义文件路径
    downloads_dir = "downloads"
    md_files = [
        ("tech_whitepaper_cn.md", "tech_whitepaper_cn.docx", "tech_whitepaper_cn.pdf"),
        ("tech_whitepaper_en.md", "tech_whitepaper_en.docx", "tech_whitepaper_en.pdf")
    ]
    
    # 确保downloads目录存在
    if not os.path.exists(downloads_dir):
        os.makedirs(downloads_dir)
    
    # 转换每个MD文件为DOCX和PDF
    for md_file, docx_file, pdf_file in md_files:
        md_path = os.path.join(downloads_dir, md_file)
        docx_path = os.path.join(downloads_dir, docx_file)
        pdf_path = os.path.join(downloads_dir, pdf_file)
        
        # 检查MD文件是否存在
        if os.path.exists(md_path):
            print(f"正在处理: {md_file}")
            
            # 读取MD文件内容
            md_content = read_md_file(md_path)
            
            # 转换为DOCX
            convert_md_to_docx(md_content, docx_path)
            
            # 转换为PDF
            convert_md_to_pdf(md_content, pdf_path)
        else:
            print(f"警告: 文件 {md_path} 不存在")

if __name__ == "__main__":
    main()