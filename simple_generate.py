import os
from docx import Document
from fpdf import FPDF

def read_md_file(file_path):
    """读取MD文件内容"""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def md_to_docx(md_content, output_path):
    """将MD内容转换为DOCX格式"""
    doc = Document()
    
    lines = md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            if line.strip():
                doc.add_paragraph(line)
    
    doc.save(output_path)
    print(f"已生成DOCX文件: {output_path}")

def md_to_pdf(md_content, output_path):
    """将MD内容转换为PDF格式"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # 尝试使用默认字体
    pdf.set_font("Arial", size=12)
    
    lines = md_content.split('\n')
    for line in lines:
        # 简单处理标题
        if line.startswith('# '):
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, line[2:], ln=True)
            pdf.set_font("Arial", size=12)
        elif line.startswith('## '):
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, line[3:], ln=True)
            pdf.set_font("Arial", size=12)
        elif line.startswith('### '):
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, line[4:], ln=True)
            pdf.set_font("Arial", size=12)
        else:
            if line.strip():
                # 处理可能的编码问题
                try:
                    pdf.cell(0, 10, line, ln=True)
                except:
                    # 如果有编码问题，尝试替换特殊字符
                    safe_line = line.encode('ascii', 'ignore').decode('ascii')
                    pdf.cell(0, 10, safe_line, ln=True)
    
    pdf.output(output_path)
    print(f"已生成PDF文件: {output_path}")

def main():
    downloads_dir = "downloads"
    
    # 确保目录存在
    if not os.path.exists(downloads_dir):
        print(f"目录 {downloads_dir} 不存在")
        return
    
    # 处理中英文技术白皮书
    files = [
        ("tech_whitepaper_cn.md", "tech_whitepaper_cn.docx", "tech_whitepaper_cn.pdf"),
        ("tech_whitepaper_en.md", "tech_whitepaper_en.docx", "tech_whitepaper_en.pdf")
    ]
    
    for md_name, docx_name, pdf_name in files:
        md_path = os.path.join(downloads_dir, md_name)
        
        if os.path.exists(md_path):
            print(f"正在处理: {md_name}")
            
            # 读取MD文件
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 生成DOCX
            docx_path = os.path.join(downloads_dir, docx_name)
            md_to_docx(content, docx_path)
            
            # 生成PDF
            pdf_path = os.path.join(downloads_dir, pdf_name)
            md_to_pdf(content, pdf_path)
        else:
            print(f"文件不存在: {md_path}")

if __name__ == "__main__":
    main()