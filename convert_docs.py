import os
from docx import Document
from fpdf import FPDF

# 检查文件是否存在
downloads_dir = "downloads"
md_files = ["tech_whitepaper_cn.md", "tech_whitepaper_en.md"]

for md_file in md_files:
    md_path = os.path.join(downloads_dir, md_file)
    if not os.path.exists(md_path):
        print(f"错误: 文件 {md_path} 不存在")
        exit(1)

print("所有MD文件都存在，开始转换...")

def read_md_file(file_path):
    """读取MD文件内容"""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def md_to_docx(md_content, output_path):
    """将MD内容转换为DOCX格式"""
    doc = Document()
    
    lines = md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            if line.strip():
                doc.add_paragraph(line)
    
    doc.save(output_path)
    print(f"已生成DOCX文件: {output_path}")

def md_to_pdf(md_content, output_path):
    """将MD内容转换为PDF格式"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # 使用Arial字体
    pdf.set_font("Arial", size=12)
    
    lines = md_content.split('\n')
    for line in lines:
        # 处理标题
        if line.startswith('# '):
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, line[2:].encode('latin-1', 'replace').decode('latin-1'), ln=True)
            pdf.set_font("Arial", size=12)
        elif line.startswith('## '):
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, line[3:].encode('latin-1', 'replace').decode('latin-1'), ln=True)
            pdf.set_font("Arial", size=12)
        elif line.startswith('### '):
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, line[4:].encode('latin-1', 'replace').decode('latin-1'), ln=True)
            pdf.set_font("Arial", size=12)
        else:
            if line.strip():
                try:
                    pdf.cell(0, 10, line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                except:
                    # 如果有编码问题，尝试替换特殊字符
                    safe_line = line.encode('ascii', 'ignore').decode('ascii')
                    pdf.cell(0, 10, safe_line, ln=True)
    
    pdf.output(output_path)
    print(f"已生成PDF文件: {output_path}")

# 处理中英文技术白皮书
files = [
    ("tech_whitepaper_cn.md", "tech_whitepaper_cn.docx", "tech_whitepaper_cn.pdf"),
    ("tech_whitepaper_en.md", "tech_whitepaper_en.docx", "tech_whitepaper_en.pdf")
]

for md_name, docx_name, pdf_name in files:
    md_path = os.path.join(downloads_dir, md_name)
    
    if os.path.exists(md_path):
        print(f"正在处理: {md_name}")
        
        # 读取MD文件
        content = read_md_file(md_path)
        
        # 生成DOCX
        docx_path = os.path.join(downloads_dir, docx_name)
        md_to_docx(content, docx_path)
        
        # 生成PDF
        pdf_path = os.path.join(downloads_dir, pdf_name)
        md_to_pdf(content, pdf_path)
    else:
        print(f"文件不存在: {md_path}")

print("所有文件转换完成!")